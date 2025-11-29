"""
File Parser Service
Extracts text content from various file formats (PDF, DOCX, TXT, MD, etc.)
"""
import io
import re
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

# Maximum file size (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024

# Supported file extensions
SUPPORTED_EXTENSIONS = {
    '.txt': 'text/plain',
    '.md': 'text/markdown',
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc': 'application/msword',
    '.rtf': 'application/rtf',
    '.html': 'text/html',
    '.htm': 'text/html',
    '.json': 'application/json',
    '.csv': 'text/csv',
}


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ''
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
    return text.strip()


def get_file_extension(filename: str) -> str:
    """Get lowercase file extension."""
    if '.' in filename:
        return '.' + filename.rsplit('.', 1)[1].lower()
    return ''


def is_supported_file(filename: str) -> bool:
    """Check if file type is supported."""
    ext = get_file_extension(filename)
    return ext in SUPPORTED_EXTENSIONS


def parse_txt_file(file_content: bytes) -> Tuple[str, str]:
    """Parse plain text file."""
    # Try different encodings
    encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
    
    for encoding in encodings:
        try:
            text = file_content.decode(encoding)
            return text, 'text'
        except UnicodeDecodeError:
            continue
    
    # Fallback with error handling
    return file_content.decode('utf-8', errors='ignore'), 'text'


def parse_markdown_file(file_content: bytes) -> Tuple[str, str]:
    """Parse markdown file."""
    text, _ = parse_txt_file(file_content)
    return text, 'markdown'


def parse_pdf_file(file_content: bytes) -> Tuple[str, str]:
    """Parse PDF file using PyPDF2."""
    try:
        import PyPDF2
        
        pdf_file = io.BytesIO(file_content)
        reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return '\n\n'.join(text_parts), 'pdf'
        
    except ImportError:
        raise ValueError('PDF parsing requires PyPDF2. Install with: pip install PyPDF2')
    except Exception as e:
        raise ValueError(f'Failed to parse PDF: {str(e)}')


def parse_docx_file(file_content: bytes) -> Tuple[str, str]:
    """Parse DOCX file using python-docx."""
    try:
        from docx import Document
        
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return '\n\n'.join(text_parts), 'docx'
        
    except ImportError:
        raise ValueError('DOCX parsing requires python-docx. Install with: pip install python-docx')
    except Exception as e:
        raise ValueError(f'Failed to parse DOCX: {str(e)}')


def parse_html_file(file_content: bytes) -> Tuple[str, str]:
    """Parse HTML file."""
    try:
        from bs4 import BeautifulSoup
        
        text, _ = parse_txt_file(file_content)
        soup = BeautifulSoup(text, 'html.parser')
        
        # Remove script and style elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header']):
            element.decompose()
        
        return soup.get_text(separator='\n'), 'html'
        
    except ImportError:
        # Fallback: strip HTML tags with regex
        text, _ = parse_txt_file(file_content)
        text = re.sub(r'<[^>]+>', '', text)
        return text, 'html'


def parse_json_file(file_content: bytes) -> Tuple[str, str]:
    """Parse JSON file and convert to readable text."""
    import json
    
    text, _ = parse_txt_file(file_content)
    
    try:
        data = json.loads(text)
        # Pretty print JSON
        formatted = json.dumps(data, indent=2, ensure_ascii=False)
        return formatted, 'json'
    except json.JSONDecodeError:
        return text, 'json'


def parse_csv_file(file_content: bytes) -> Tuple[str, str]:
    """Parse CSV file and convert to readable text."""
    import csv
    
    text, _ = parse_txt_file(file_content)
    
    try:
        reader = csv.reader(io.StringIO(text))
        rows = []
        for row in reader:
            rows.append(' | '.join(row))
        return '\n'.join(rows), 'csv'
    except Exception:
        return text, 'csv'


def parse_file(filename: str, file_content: bytes) -> dict:
    """
    Parse a file and extract its text content.
    
    Args:
        filename: Original filename with extension
        file_content: Raw file bytes
        
    Returns:
        dict with keys: success, content, file_type, error
    """
    try:
        # Check file size
        if len(file_content) > MAX_FILE_SIZE:
            return {
                'success': False,
                'error': f'File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB'
            }
        
        # Get file extension
        ext = get_file_extension(filename)
        
        if not ext:
            return {
                'success': False,
                'error': 'Could not determine file type. Please ensure the file has an extension.'
            }
        
        if ext not in SUPPORTED_EXTENSIONS:
            return {
                'success': False,
                'error': f'Unsupported file type: {ext}. Supported types: {", ".join(SUPPORTED_EXTENSIONS.keys())}'
            }
        
        # Parse based on extension
        if ext in ['.txt']:
            content, file_type = parse_txt_file(file_content)
        elif ext in ['.md']:
            content, file_type = parse_markdown_file(file_content)
        elif ext in ['.pdf']:
            content, file_type = parse_pdf_file(file_content)
        elif ext in ['.docx']:
            content, file_type = parse_docx_file(file_content)
        elif ext in ['.doc']:
            # .doc files are harder to parse, try as binary text
            return {
                'success': False,
                'error': 'Legacy .doc format is not supported. Please convert to .docx'
            }
        elif ext in ['.html', '.htm']:
            content, file_type = parse_html_file(file_content)
        elif ext in ['.json']:
            content, file_type = parse_json_file(file_content)
        elif ext in ['.csv']:
            content, file_type = parse_csv_file(file_content)
        elif ext in ['.rtf']:
            # RTF is complex, try basic text extraction
            content, file_type = parse_txt_file(file_content)
            file_type = 'rtf'
        else:
            content, file_type = parse_txt_file(file_content)
        
        # Clean the content
        content = clean_text(content)
        
        if not content or len(content) < 10:
            return {
                'success': False,
                'error': 'Could not extract meaningful content from the file'
            }
        
        return {
            'success': True,
            'content': content,
            'file_type': file_type,
            'filename': filename
        }
        
    except ValueError as e:
        return {
            'success': False,
            'error': str(e)
        }
    except Exception as e:
        logger.error(f"Error parsing file {filename}: {str(e)}")
        return {
            'success': False,
            'error': f'Failed to parse file: {str(e)}'
        }
